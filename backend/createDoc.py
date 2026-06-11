from docx import Document


class Doc: 
    """
    Class will create a resume that is in document format.
    """
    def __init__(self):
        self.document = Document()

    
    def txt2Doc(self, resume: str) -> Document:
        """
        Method rconverts resume.txt to a docment format 

        Args: 
            resume (str): Is the resume in txt format.
        Returns:
            Docuent: Will return resume in a document format.
        """
        with open(resume) as file:
            for line in file:
                if line.startswith("<h1>"):
                    self.document.add_heading(line.replace("<h1>", "").replace("</h1>", "").strip(), 1)                     
                else:
                    self.document.add_paragraph(line.strip())

        self.document.save("docs/demo.docx")
