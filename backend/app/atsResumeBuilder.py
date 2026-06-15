from openai import OpenAI
from pathlib import Path
from docx import Document
import os
import sys
import datetime as dt
from pathlib import Path


class ATSResume:
    """
    Builds an ATS friendly resume using AI.
    """
    def __init__(self, resume):
        """
        Args:
            resume (str): the path of where the resume is located.
        """
        self.document = Document()
        self.resume = Path(resume)


    def readResume(self):
        # String is immutable so you are creating a new string everytime you add to.
        # You are using a lot of memory 
        """ 
            Reads the resum and adds it to a string. 
        """
        self.content_of_resume = ""
        with open(self.resume, "r") as file:
            for line in file:
                self.content_of_resume += line + "\n"


    def makeATSResume(self):
        """
        Creates a ATS friendly resume using  OPENAI
        """
        client = OpenAI(
        api_key=os.environ.get("GROQ_API_KEY"),
        base_url="https://api.groq.com/openai/v1",
        )

        prompt = f"""
        Task: Using the {self.content_of_resume} and the {self.description} , generate a tailored, ATS-friendly CV in plain text format.
        Instructions:
        Keyword Analysis: Identify the top 10 most important technical skills, tools, and soft skills mentioned in the {self.description}  and {self.requirements}. Ensure these keywords appear naturally in the "Technical Skills" and "Professional Experience" sections.
        Impact-First Bullet Points: Rewrite existing experience into the "Action Verb + Task + Result" formula. If the Master CV {self.resume} mentions "Developed a bank app," and the {self.description} emphasizes "database optimization,"

        Professional Summary: Write a 3-line summary that mirrors the "Years of Experience" and "Primary Tech Stack" requested in the {self.description}.
        ATS Formatting Rules:
        Use standard headers (Professional Summary, Technical Skills, Experience, Education).
        Avoid tables, columns, or graphics.
        Use simple bullet points.

        Constraint: Do not hallucinate or invent experiences not found in the Master_CV {self.resume}. If a required skill in the {self.description} is missing from the Master CV, do not add it; instead, emphasize related transferable skills.
        Do not use emoji's.
        Input Data:
        Master CV: {self.content_of_resume}
        Job Description: {self.description}
"""

        response = client.responses.create(
            input=prompt,
            instructions= "Role: You are an expert Technical Career Coach and ATS (Applicant Tracking System) Optimization Specialist. Your goal is to rewrite a candidate's Master CV to perfectly align with a specific Job Description while maintaining 100percent honesty and professional integrity. Only respond with the CV. When writing cv write it in html tags. ",
            model="openai/gpt-oss-20b",
        )
        self.newResume = response.output_text

    def writeToFile(self):
        """
        loads content into file and save the file in the root directory.
        """
        self.APP_DIR = Path(__file__).resolve().parent
        self.BACKEND_DIR = self.APP_DIR.parent

        self.date = str(dt.datetime.today()).split(" ")
        
        logs_dir =  f"{self.BACKEND_DIR}/resumes/.logs/{self.date[0]}"
        resume_dir =  f"{self.BACKEND_DIR}/resumes/{self.date[0]}"

        directory = f"/resumes/.logs/{self.date[0]}/"

        if os.path.isdir(logs_dir) !=  True :
            os.makedirs(logs_dir, exist_ok=True)
            os.makedirs(resume_dir, exist_ok=True)

        with open(f"{logs_dir}/{self.companyName}-{self.position}.txt", "w") as file:
            for line in self.newResume:
                file.writelines(line)
        
        resume = f"{logs_dir}/{self.companyName}-{self.position}.txt"
        self.txt2Doc(resume)

    def txt2Doc(self, resume: str) -> Document:
        """
        Method rconverts resume.txt to a docment format 

        Args: 
            resume (str): Is the resume in txt format.
        Returns:
            Docuent: Will return resume in a document format.
        """
        with open(resume) as file:
            for line in file:
                if line.startswith("<h1>"):
                    self.document.add_heading(line.replace("<h1>", "").replace("</h1>", "").strip(), 1)                     
                elif line.startswith("<h2>"):
                    self.document.add_heading(line.replace("<h2>","").replace("</h2>","").strip(), 1)
                elif line.startswith("<p>"):
                    self.document.add_paragraph(line.replace("<p>", "").replace("</p>", "").strip())
                elif line.startswith("<li>"):
                    self.document.add_paragraph(line.replace("<li>", "-").replace("</li>", "").strip())
                else:
                    continue
        
        self.document.save(f"{self.BACKEND_DIR}/resumes/{self.date[0]}/{self.companyName}-{self.position}.docx")

    def jobListingInfo(self):
        """
        Gets Job Listing information.
        """
        self.companyName = input("\tCompany Name\n\t> ")
        self.position = input("\tPosition\n\t> ")
        self.description = input("\tDescription\n\t> ")
        self.requirements = input("\tRequirements\n\t> ")

